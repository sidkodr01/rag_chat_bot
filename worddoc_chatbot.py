import os
from langchain.text_splitter import CharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
import numpy as np
from langchain.memory import ConversationBufferMemory
from docx import Document as DocxDocument  # For reading Word documents
from langchain.docstore.document import Document  # LangChain Document class

# Configuration variables
CHUNK_SIZE = 300
CHUNK_OVERLAP = 50
MAX_TOKENS = 15000
MODEL_NAME = "gemini-pro"  # Google's model name
TEMPERATURE = 0.4

# Set up Google API key
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    GOOGLE_API_KEY = input("Please enter your Google API key: ")
    os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY

def read_word_document(file_path):
    """
    Reads a Word document and extracts its text content.
    """
    try:
        doc = DocxDocument(file_path)
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        return content
    except Exception as e:
        print(f"Error reading the Word document: {e}")
        return []

def clean_content(content_list):
    """
    Cleans the content by removing very short or unwanted items.
    """
    cleaned = [text for text in content_list if len(text) > 20 and not any(item in text.lower() for item in ['sign up', 'sign in', 'cookie', 'privacy policy'])]
    return cleaned

def process_word_document(file_path):
    """
    Processes a Word document and splits it into chunks.
    """
    content = read_word_document(file_path)
    if not content:
        raise ValueError("No content could be read from the Word document.")
    
    # Convert the content into LangChain Document objects
    documents = [Document(page_content=text, metadata={"source": file_path}) for text in content]
    
    print(f"\nNumber of documents loaded: {len(documents)}")
    if documents:
        print("Sample of loaded content:")
        print(documents[0].page_content[:200] + "...")
    
    # Split the documents into chunks
    text_splitter = CharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    texts = text_splitter.split_documents(documents)
    print(f"Number of text chunks after splitting: {len(texts)}")
    return texts

def print_sample_embeddings(texts, embeddings):
    """
    Prints a sample of the embeddings for debugging.
    """
    if texts:
        sample_text = texts[0].page_content
        sample_embedding = embeddings.embed_query(sample_text)
        print("\nSample Text:")
        print(sample_text[:200] + "..." if len(sample_text) > 200 else sample_text)
        print("\nSample Embedding (first 10 dimensions):")
        print(np.array(sample_embedding[:10]))
        print(f"\nEmbedding shape: {np.array(sample_embedding).shape}")
    else:
        print("No texts available for embedding sample.")

# Set up Google language model
llm = ChatGoogleGenerativeAI(
    model=MODEL_NAME,
    temperature=TEMPERATURE,
    max_tokens=MAX_TOKENS
)

# Set up the retrieval-based QA system with a simplified prompt template
template = """Context: {context}

Question: {question}

Answer the question concisely based only on the given context. If the context doesn't contain relevant information, say "I don't have enough information to answer that question."

But, if the question is generic, then go ahead and answer the question, example what is a electric vehicle?
"""

PROMPT = PromptTemplate(
    template=template, input_variables=["context", "question"]
)

memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

def rag_pipeline(query, qa_chain, vectorstore):
    """
    Runs the RAG pipeline to answer a query.
    """
    relevant_docs = vectorstore.similarity_search_with_score(query, k=3)
    
    print("\nTop 3 most relevant chunks:")
    context = ""
    for i, (doc, score) in enumerate(relevant_docs, 1):
        print(f"{i}. Relevance Score: {score:.4f}")
        print(f"   Content: {doc.page_content[:200]}...")
        print()
        context += doc.page_content + "\n\n"

    # Print the full prompt
    full_prompt = PROMPT.format(context=context, question=query)
    print("\nFull Prompt sent to the model:")
    print(full_prompt)
    print("\n" + "="*50 + "\n")

    response = qa_chain.invoke({"query": query})
    return response['result']

if __name__ == "__main__":
    print("Welcome to the Enhanced RAG Pipeline for Word Documents.")
    
    while True:
        file_path = input("Please enter the path to the Word document you want to query (or 'quit' to exit): ")
        if file_path.lower() == 'quit':
            print("Exiting the program. Goodbye!")
            break
        
        if not os.path.exists(file_path):
            print("The specified file does not exist. Please try again.")
            continue
        
        try:
            print("Processing Word document content...")
            texts = process_word_document(file_path)
            
            if not texts:
                print("No content found in the Word document. Please try a different file.")
                continue
            
            print("Creating embeddings and vector store...")
            embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")  # Specify the model here
            
            print_sample_embeddings(texts, embeddings)
            
            vectorstore = FAISS.from_documents(texts, embeddings)
            
            qa = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=vectorstore.as_retriever(),
                memory=memory,
                chain_type_kwargs={"prompt": PROMPT}
            )
            
            print("\nRAG Pipeline initialized. You can now enter your queries.")
            print("Enter 'new' to query a new document or 'quit' to exit the program.")
            
            while True:
                user_query = input("\nEnter your query: ")
                if user_query.lower() == 'quit':
                    print("Exiting the program. Goodbye!")
                    exit()
                elif user_query.lower() == 'new':
                    break
                
                result = rag_pipeline(user_query, qa, vectorstore)
                print(f"RAG Response: {result}")
        
        except Exception as e:
            print(f"An error occurred: {e}")
            print("Please try a different document or check the file format.")